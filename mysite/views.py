# -*- encoding: utf-8 -*-

from django import template
from django.contrib.auth.decorators import login_required
from django.http import HttpResponse, HttpResponseRedirect, FileResponse
from django.shortcuts import render
from django.template import loader
from django.urls import reverse
from docx import settings

from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.platypus import Table, TableStyle, Paragraph
from reportlab.pdfgen.canvas import Canvas
from reportlab.lib.styles import getSampleStyleSheet

from docx import *
from docx.shared import Inches

from .forms import WendlerForm
from reportlab.pdfgen import canvas
import io


'''@login_required(login_url="/login/")
def wendler_view(request):
    context = {'segment': 'wendler',
               'number':3}
    html_template = loader.get_template('home/wendler.html')
    return HttpResponse(html_template.render(context, request))'''

global_wendler_list = {}

def some_view(request):

    styles = getSampleStyleSheet()
    style = styles["BodyText"]

    buffer = io.BytesIO()

    canv = canvas.Canvas(buffer)

    empty_list = []

    empty_list.append(['Week No.', 'Set 1', "Set 2", "Set 3", "Set 4"])
    empty_list.append(['Week 1'] + list(global_wendler_list['Week 1'].values()))
    empty_list.append(['Week 2'] + list(global_wendler_list['Week 2'].values()))
    empty_list.append(['Week 3'] + list(global_wendler_list['Week 3'].values()))
    empty_list.append(['Week 4'] + list(global_wendler_list['Week 4'].values()))

    print(empty_list)

    header = Paragraph("<bold><font size=18>Wendler Exercise List</font></bold>", style)

    data = empty_list
    t = Table(data)
    t.setStyle(TableStyle([("BOX", (0, 0), (-1, -1), 0.25, colors.black),
                           ('INNERGRID', (0, 0), (-1, -1), 0.25, colors.black)]))
    data_len = len(data)

    for each in range(data_len):
        if each % 2 == 0:
            bg_color = colors.whitesmoke
        else:
            bg_color = colors.lightgrey

        t.setStyle(TableStyle([('BACKGROUND', (0, each), (-1, each), bg_color)]))

    aW = 540
    aH = 720

    w, h = header.wrap(aW, aH)
    header.drawOn(canv, 72, 800)
    aH = aH - h
    w, h = t.wrap(aW, aH)
    t.drawOn(canv, 72, aH - h)
    canv.save()


    # FileResponse sets the Content-Disposition header so that browsers
    # present the option to save the file.
    buffer.seek(0)
    return FileResponse(buffer, as_attachment=True, filename='WendlerSheet.pdf')

def word_doc_view(request):

    document = Document()
    docx_title="WendlerSheet.docx"
    # ---- Cover Letter ----
    document.add_paragraph("Wendler Exercise List")

    document.add_paragraph('Week 1' + str(global_wendler_list['Week 1'])[1:-1])

    document.add_paragraph('Week 2' + str(global_wendler_list['Week 2'])[1:-1])

    document.add_paragraph('Week 3' + str(global_wendler_list['Week 3'])[1:-1])

    document.add_paragraph('Week 4' + str(global_wendler_list['Week 4'])[1:-1])


    document.add_page_break()

    # Prepare document for download
    # -----------------------------
    f = io.BytesIO()
    document.save(f)
    length = f.tell()
    f.seek(0)
    response = HttpResponse(
        f.getvalue(),
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    response['Content-Disposition'] = 'attachment; filename=' + docx_title
    response['Content-Length'] = length
    return response

def wendler_view(request):
    if request.method == 'POST':

        form = WendlerForm(request.POST)

        if form.is_valid():
            # Takes the input of one Rep Max and assigns it to the number variable
            number = int(request.POST['weight'])
            global global_wendler_list

            # The list of percentages from the Wendler 531 regimen
            percentage_list = [0.40, 0.50, 0.60, 0.65, 0.70, 0.75, 0.80, 0.85, 0.90, 0.95]

            # Initializes a dictionary for containing calculated exercise weights
            calculated_dict = {}

            # Loop for iterating over percentage list and number to propagate calculated_dict
            for num in percentage_list:

                # Multiplies One Rep Max with Percentage then Subtracts empty bar weight and Divides by two for both sides
                calc_num = (number * num - 20) / 2

                # Conditional for rounding down to the nearest 2.5 kg (smallest plate)
                if (calc_num % 2.5) < 1.25:
                    calc_num = calc_num - (calc_num % 2.5)

                # Conditional for rounding up to the nearest 2.5 kg (smallest plate)
                else:
                    calc_num = calc_num + 2.5 - (calc_num % 2.5)

                # Conditional for replacing negatives with zeros
                if calc_num < 0:
                    calc_num = 0

                # Assigning weight value to the percentage key
                calculated_dict[num] = calc_num

            # Initializes a dictionary for containing calculated exercise sets
            exercise_dict = {
                'Week 1': {'Set 1': str(calculated_dict[0.4]) + 'kgx5',
                          'Set 2': str(calculated_dict[0.65]) + 'kgx5',
                          'Set 3': str(calculated_dict[0.75]) + 'kgx5',
                          'Set 4': str(calculated_dict[0.85]) + 'kgx5'},

                'Week 2': {'Set 1': str(calculated_dict[0.4]) + 'kgx3',
                          'Set 2': str(calculated_dict[0.7]) + 'kgx3',
                          'Set 3': str(calculated_dict[0.8]) + 'kgx3',
                          'Set 4': str(calculated_dict[0.9]) + 'kgx3'},

                'Week 3': {'Set 1': str(calculated_dict[0.4]) + 'kgx5',
                          'Set 2': str(calculated_dict[0.75]) + 'kgx5',
                          'Set 3': str(calculated_dict[0.85]) + 'kgx3',
                          'Set 4': str(calculated_dict[0.95]) + 'kgx1'},

                'Week 4': {'Set 1': str(calculated_dict[0.4]) + 'kgx5',
                          'Set 2': str(calculated_dict[0.4]) + 'kgx5',
                          'Set 3': str(calculated_dict[0.5]) + 'kgx5',
                          'Set 4': str(calculated_dict[0.6]) + 'kgx5'},
            }

            print("Woop")

            global_wendler_list = exercise_dict

            return render(request, 'wendler.html',
                          {'form': form, 'number': number, 'calculated_dict': exercise_dict})

    else:
        form = WendlerForm()

    return render(request, 'wendler.html', {'form': form})


def pages(request):
    context = {}
    # All resource paths end in .html.
    # Pick out the html file name from the url. And load that template.
    try:

        load_template = request.path.split('/')[-1]

        if load_template == 'admin':
            return HttpResponseRedirect(reverse('admin:wendler'))
        context['segment'] = load_template
        html_template = loader.get_template('home/' + load_template)
        return HttpResponse(html_template.render(context, request))

    except template.TemplateDoesNotExist:

        html_template = loader.get_template('home/page-404.html')
        return HttpResponse(html_template.render(context, request))

    except:
        html_template = loader.get_template('home/page-500.html')
        return HttpResponse(html_template.render(context, request))
